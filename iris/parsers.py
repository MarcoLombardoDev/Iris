# -*- coding: utf-8 -*-
"""Extraction of (company, email) pairs from the supported documents.

The module never imports Tkinter: it is pure logic, testable without a
display.
"""

import csv
import os
import re
from typing import Callable, Iterable, List, NamedTuple, Optional, Sequence

from .i18n import t

#: File extensions handled by the application.
SUPPORTED_EXTENSIONS = (".pdf", ".xlsx", ".xlsm", ".xls", ".docx", ".txt", ".csv")

#: Email address pattern.
#: NB: the pre-2.0 pattern used ``[A-Z|a-z]{2,}`` and wrongly accepted a
#: pipe character inside the top-level domain.
EMAIL_REGEX = re.compile(
    r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9](?:[A-Za-z0-9\-]*[A-Za-z0-9])?"
    r"(?:\.[A-Za-z0-9](?:[A-Za-z0-9\-]*[A-Za-z0-9])?)*\.[A-Za-z]{2,}"
)

#: Separators recognised in "Company <sep> email" lines.
_SEPARATORS = (";", ",", "\t", "|")

#: Surrounding characters stripped from a company name.
#: NB: a trailing dot is kept on purpose — it belongs to legal forms such as
#: "S.r.l." or "Inc.".
_COMPANY_CLEAN_RE = re.compile(r"^[\s\-–—:;,.*•>]+|[\s\-–—:;,*•<]+$")

Logger = Optional[Callable[[str], None]]


class Recipient(NamedTuple):
    """A company/email pair extracted from a document."""

    company: str
    email: str


class UnsupportedFormatError(ValueError):
    """The file format is not handled, or the required library is missing."""


def is_valid_email(value: str) -> bool:
    """True when the string is a syntactically valid email address."""
    if not value:
        return False
    return EMAIL_REGEX.fullmatch(value.strip()) is not None


def normalize_email(value: str) -> str:
    """Clean an email address (spaces, ``mailto:``, angle brackets)."""
    cleaned = (value or "").strip().strip("<>").strip()
    if cleaned.lower().startswith("mailto:"):
        cleaned = cleaned[7:]
    return cleaned


def clean_company(value: str) -> str:
    """Normalise a company name, trimming surrounding punctuation."""
    if not value:
        return ""
    cleaned = re.sub(r"\s+", " ", str(value)).strip()
    cleaned = _COMPANY_CLEAN_RE.sub("", cleaned)
    return cleaned.strip()


def fallback_company(email: str) -> str:
    """Company name derived from the address domain, used as a last resort."""
    try:
        domain = email.split("@", 1)[1].split(".")[0]
    except IndexError:
        return t("parsers.fallback_company_generic")
    if not domain:
        return t("parsers.fallback_company_generic")
    return t("parsers.fallback_company", domain=domain.capitalize())


def dedupe(recipients: Iterable[Recipient]) -> List[Recipient]:
    """Drop duplicated addresses, keeping the first occurrence."""
    seen = set()
    unique: List[Recipient] = []
    for company, email in recipients:
        key = email.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(Recipient(company, email))
    return unique


def _make_recipient(company: str, email: str) -> Optional[Recipient]:
    """Build a ``Recipient``, validating the address and cleaning the name."""
    email = normalize_email(email)
    if not is_valid_email(email):
        return None
    company = clean_company(company)
    if not company:
        company = fallback_company(email)
    return Recipient(company[:120], email)


# --------------------------------------------------------------------------
# Plain text (txt, docx, PDF fallback)
# --------------------------------------------------------------------------

def extract_from_lines(lines: Sequence[str]) -> List[Recipient]:
    """Extract company/email pairs from a list of text lines.

    Rules, in order:

    1. line with a separator (``;`` ``,`` tab ``|``): the first part holding an
       address is the email, the first non-email part is the company name;
    2. line without a separator but with an address: the text preceding the
       address is the company name;
    3. no name found: the address domain is used instead.
    """
    results: List[Recipient] = []
    for raw_line in lines:
        line = (raw_line or "").strip()
        if not line:
            continue

        match = EMAIL_REGEX.search(line)
        if not match:
            continue

        separator = next((sep for sep in _SEPARATORS if sep in line), None)
        company = ""
        email = match.group(0)

        if separator:
            parts = [part.strip() for part in line.split(separator) if part.strip()]
            email_part_index = None
            for index, part in enumerate(parts):
                part_match = EMAIL_REGEX.search(part)
                if part_match:
                    email = part_match.group(0)
                    email_part_index = index
                    break
            for index, part in enumerate(parts):
                if index == email_part_index:
                    continue
                if EMAIL_REGEX.search(part):
                    continue
                if clean_company(part):
                    company = part
                    break
        else:
            company = line[: match.start()]

        recipient = _make_recipient(company, email)
        if recipient:
            results.append(recipient)
    return results


def extract_from_text(text: str) -> List[Recipient]:
    """Extract company/email pairs from a block of text."""
    results = extract_from_lines((text or "").splitlines())
    if results:
        return dedupe(results)

    # No pair recognised: at least collect the addresses present.
    fallback = []
    for match in EMAIL_REGEX.finditer(text or ""):
        recipient = _make_recipient("", match.group(0))
        if recipient:
            fallback.append(recipient)
    return dedupe(fallback)


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def extract_from_pdf_text(text: str) -> List[Recipient]:
    """Extract pairs from text coming out of a PDF.

    PDFs often lack separators, so the context preceding the address is used:
    same line first, then the previous lines.
    """
    lines = (text or "").splitlines()
    results: List[Recipient] = []

    for index, line in enumerate(lines):
        for match in EMAIL_REGEX.finditer(line):
            email = match.group(0)
            before = line[: match.start()]

            separator = next((sep for sep in _SEPARATORS if sep in before), None)
            if separator:
                before = before.rsplit(separator, 1)[0]

            company = clean_company(before)
            if not company and index > 0:
                # Walk backwards to the first non-empty line without an address.
                for previous in reversed(lines[max(0, index - 3):index]):
                    if EMAIL_REGEX.search(previous):
                        continue
                    candidate = clean_company(previous)
                    if candidate:
                        company = candidate
                        break

            recipient = _make_recipient(company, email)
            if recipient:
                results.append(recipient)
    return dedupe(results)


def extract_from_pdf(path: str, log: Logger = None) -> List[Recipient]:
    """Extract company/email pairs from a PDF file (requires PyMuPDF)."""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover - depends on the environment
        raise UnsupportedFormatError(t("parsers.missing_pymupdf")) from exc

    results: List[Recipient] = []
    doc = fitz.open(path)
    try:
        for page_number, page in enumerate(doc, start=1):
            if log:
                log(t("log.analysis_page", page=page_number))
            results.extend(extract_from_pdf_text(page.get_text()))
    finally:
        doc.close()
    return dedupe(results)


# --------------------------------------------------------------------------
# Spreadsheets
# --------------------------------------------------------------------------

def extract_from_rows(rows: Iterable[Sequence[object]]) -> List[Recipient]:
    """Extract pairs from tabular rows.

    Column order is not assumed: for every row the first cell holding a valid
    address is the email, and the first non-empty cell that is not the address
    is the company name. A header row is dropped automatically because it
    holds no valid address.
    """
    results: List[Recipient] = []
    for row in rows:
        cells = ["" if cell is None else str(cell).strip() for cell in row]
        if not any(cells):
            continue

        email = ""
        email_index = None
        for index, cell in enumerate(cells):
            candidate = normalize_email(cell)
            if is_valid_email(candidate):
                email = candidate
                email_index = index
                break
            match = EMAIL_REGEX.search(cell)
            if match:
                email = match.group(0)
                email_index = index
                break
        if not email:
            continue

        company = ""
        for index, cell in enumerate(cells):
            if index == email_index or not cell:
                continue
            if EMAIL_REGEX.search(cell):
                continue
            company = cell
            break

        recipient = _make_recipient(company, email)
        if recipient:
            results.append(recipient)
    return dedupe(results)


def extract_from_xlsx(path: str) -> List[Recipient]:
    """Extract pairs from a modern Excel file (.xlsx/.xlsm)."""
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover - depends on the environment
        raise UnsupportedFormatError(t("parsers.missing_openpyxl")) from exc

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    try:
        results: List[Recipient] = []
        for sheet in workbook.worksheets:
            results.extend(extract_from_rows(sheet.iter_rows(values_only=True)))
        return dedupe(results)
    finally:
        workbook.close()


def extract_from_xls(path: str) -> List[Recipient]:
    """Extract pairs from a legacy Excel file (.xls).

    ``openpyxl`` does not support the binary .xls format: ``xlrd`` is required.
    """
    try:
        import xlrd
    except ImportError as exc:  # pragma: no cover - depends on the environment
        raise UnsupportedFormatError(t("parsers.missing_xlrd")) from exc

    book = xlrd.open_workbook(path)
    try:
        results: List[Recipient] = []
        for sheet in book.sheets():
            rows = (sheet.row_values(index) for index in range(sheet.nrows))
            results.extend(extract_from_rows(rows))
        return dedupe(results)
    finally:
        try:
            book.release_resources()
        except Exception:
            pass


def extract_from_csv(path: str) -> List[Recipient]:
    """Extract pairs from a CSV file (delimiter detected automatically)."""
    text = _read_text_file(path)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=";,\t|")
        delimiter = dialect.delimiter
    except Exception:
        delimiter = ";" if sample.count(";") >= sample.count(",") else ","
    rows = list(csv.reader(text.splitlines(), delimiter=delimiter))
    return extract_from_rows(rows)


# --------------------------------------------------------------------------
# Word and plain text
# --------------------------------------------------------------------------

def extract_from_docx(path: str) -> List[Recipient]:
    """Extract pairs from a Word document (paragraphs and tables)."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - depends on the environment
        raise UnsupportedFormatError(t("parsers.missing_docx")) from exc

    document = docx.Document(path)
    results: List[Recipient] = []

    # Tables are structured data: treat them like spreadsheet rows.
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        results.extend(extract_from_rows(rows))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    results.extend(extract_from_lines(paragraphs))

    if not results:
        results.extend(extract_from_text("\n".join(paragraphs)))
    return dedupe(results)


def _read_text_file(path: str) -> str:
    """Read a text file trying the most common encodings."""
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as handle:
                return handle.read()
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        return handle.read()


def extract_from_txt(path: str) -> List[Recipient]:
    """Extract pairs from a plain text file."""
    return extract_from_text(_read_text_file(path))


# --------------------------------------------------------------------------
# Dispatcher
# --------------------------------------------------------------------------

def extract_from_file(path: str, log: Logger = None) -> List[Recipient]:
    """Extract company/email pairs from the given file, based on its extension.

    Raises :class:`UnsupportedFormatError` when the format is not handled or
    the required library is missing, and :class:`FileNotFoundError` when the
    file does not exist.
    """
    if not path:
        raise FileNotFoundError(t("parsers.no_file"))
    if not os.path.exists(path):
        raise FileNotFoundError(t("parsers.file_missing", path=path))

    extension = os.path.splitext(path)[1].lower()
    if extension == ".pdf":
        return extract_from_pdf(path, log=log)
    if extension in (".xlsx", ".xlsm"):
        return extract_from_xlsx(path)
    if extension == ".xls":
        return extract_from_xls(path)
    if extension == ".docx":
        return extract_from_docx(path)
    if extension == ".csv":
        return extract_from_csv(path)
    if extension == ".txt":
        return extract_from_txt(path)
    raise UnsupportedFormatError(
        t(
            "parsers.unsupported",
            ext=extension or t("parsers.no_extension"),
            formats=", ".join(SUPPORTED_EXTENSIONS),
        )
    )
